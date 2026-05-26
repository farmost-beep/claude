#!/usr/bin/env python3
"""从 markdown 生成民建社情民意 .doc 文件"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re
import os

def set_font(run, font_name, size_pt, bold=False, color=None):
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.name = font_name
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_para(doc, text, font_name, size_pt, bold=False, align=None, color=None, space_after=Pt(6), first_line_indent=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    if first_line_indent:
        p.paragraph_format.first_line_indent = first_line_indent
    run = p.add_run(text)
    set_font(run, font_name, size_pt, bold, color=color)
    return p

def clean(text):
    """清理 markdown 标记"""
    # 去掉 ** 加粗（已在字体中处理）
    return text

def parse_and_build(md_path, docx_path):
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = '仿宋'
    style.font.size = Pt(14)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # 主标题 #
        if stripped.startswith('# ') and not stripped.startswith('## '):
            add_para(doc, stripped[2:], '方正小标宋简体', 22, bold=True,
                     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(16))

        # 二级标题 ##
        elif stripped.startswith('## ') and not stripped.startswith('### '):
            text = stripped[3:]
            # 去掉 markdown 粗体标记
            text = text.replace('**', '')
            add_para(doc, text, '黑体', 15, bold=True, space_after=Pt(10))

        # 三级标题 ###
        elif stripped.startswith('### '):
            text = stripped[4:]
            text = text.replace('**', '')
            add_para(doc, text, '楷体', 14, bold=True, space_after=Pt(8))

        # 引用块 >
        elif stripped.startswith('> '):
            text = stripped[2:]
            text = text.replace('**', '')
            if '核心观点' in text:
                add_para(doc, text, '楷体', 14, bold=True, color=(0x1a, 0x56, 0xdb),
                         first_line_indent=Cm(0.74))
            else:
                add_para(doc, text, '楷体', 13, color=(0x66, 0x66, 0x66))

        # 分隔线
        elif stripped == '---':
            pass

        # 代码块
        elif stripped == '```' or stripped.startswith('```'):
            pass

        # 表格行
        elif stripped.startswith('|') and '---' not in stripped:
            cells = [c.strip().replace('**', '') for c in stripped.split('|')[1:-1]]
            text = '  |  '.join(cells)
            add_para(doc, text, '仿宋', 12, space_after=Pt(2))

        # 编号列表项 (1. 2. 3. 4.)
        elif re.match(r'^\d+\.\s+', stripped):
            text = stripped
            text = text.replace('**', '')
            add_para(doc, text, '仿宋', 14, first_line_indent=Cm(0.74))

        # 普通段落
        else:
            text = stripped.replace('**', '')
            # 跳过纯标记行
            if text in ['---', '```']:
                i += 1
                continue
            add_para(doc, text, '仿宋', 14, first_line_indent=Cm(0.74))

        i += 1

    doc.save(docx_path)
    size_kb = os.path.getsize(docx_path) / 1024
    print(f"已生成: {docx_path} ({size_kb:.1f} KB)")

if __name__ == '__main__':
    md_path = '/Users/cyingfang/claude/deliverables/career/民建社情民意-并购贷款与科创金融生态.md'
    docx_path = '/Users/cyingfang/claude/deliverables/career/民建社情民意-并购贷款与科创金融生态.doc'
    parse_and_build(md_path, docx_path)
